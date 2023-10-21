# Various methods to clean textual data

import os
import fnmatch
from docx import Document # !pip install python-docx
# from docx.core import CoreProperties
import re
import chardet # !pip install chardet
# import PyPDF2
from datetime import datetime


### Working with MS Word docs
def find_docx_files(root_dir):
    """Find all the .docx files in a folder and its subfolders

    Args:
        root_dir (str): root folder location

    Returns:
        list[str]: list of absolute file path names
    """
    docx_files = []

    for root, dirs, files in os.walk(root_dir):
        for file in files:
            if fnmatch.fnmatch(file, "*.docx"):
                docx_files.append(os.path.join(root, file))

    return docx_files

def example_find_docx_files():
    start_directory = r"C:\\Users\donnp\\OneDrive\\8. Poly"  # Replace with the path to the directory you want to search
    docx_files = find_docx_files(start_directory)

    if docx_files:
        print("Found .docx files:")
        for file in docx_files:
            print(file)
    else:
        print("No .docx files found in the specified directory and its subdirectories.")
        
def get_docx_text(file_path: str) -> None:
    """Returns the text contained in a .docx file without performing any cleaning.

    Args:
        file_path (str): absolute file path and name
    """
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
        
    # print(text)
    return text
        
def example_get_docx_text():
    file_path = "C:\\Users\donnp\\OneDrive\\8. Poly\How far does our responsibility stretch when considering metamours.docx"
    print(get_docx_text(file_path))
        
# A note on converting to json:
# When you convert text to JSON and then write it to a file using UTF-8 encoding, certain characters, especially special or non-ASCII characters, will be represented as Unicode escape sequences (e.g., "\u2014"). This is the standard way to encode such characters in JSON.
# For example, the character "—" (an em dash) is represented as "\u2014" in JSON because it's a Unicode character, and JSON uses this escape sequence format to represent characters that are not part of the basic ASCII character set.   
 
### Working with plain text docs    
def detect_encoding(file_path):
    with open(file_path, 'rb') as file:
        rawdata = file.read()
        result = chardet.detect(rawdata)
        encoding = result['encoding']
        confidence = result['confidence']

        return encoding, confidence

# Example usage
def example_detect_encoding():
    file_path = './aiutils/sample_text.txt'
    detected_encoding, confidence = detect_encoding(file_path)
    print(f"Detected encoding: {detected_encoding} with confidence: {confidence}")

def find_non_displayable_characters(file_path: str):
    """decode and reencode a file to check if there are any encoding issues

    Args:
        file_path (str): location of the file
    """
    with open(file_path, 'rb') as file:
        content = file.read()
        try:
            decoded_content = content.decode('utf-8')
            reencoded_content = decoded_content.encode('utf-8')
        except UnicodeDecodeError as e:
            # Handle the decoding error and print information about the problematic character
            print(f"UnicodeDecodeError: {e}")
            problem_character = content[e.start:e.end]
            print(f"Problematic character: {problem_character}")
            return

    print("NO non-displayable characters found.")
    
def example_find_non_displayable_characters():
    file_path = './aiutils/sample_text.txt'
    find_non_displayable_characters(file_path)

def remove_non_printable_except_newline(text):
    # Define a regex pattern to match non-printable characters
    non_printable_pattern = re.compile(r'[^\x20-\x7E\n]+')

    # Use the pattern to replace non-printable characters with an empty string
    cleaned_text = non_printable_pattern.sub('', text)

    return cleaned_text

def example_remove_non_printable_except_newline():
    text_with_non_printable = "This is a text\x07with non-printable\x1Bcharacters."
    cleaned_text = remove_non_printable_except_newline(text_with_non_printable)
    print(cleaned_text)

def clean_consecutive_spaces(text):
    # Use a regular expression to replace consecutive spaces with a single space
    cleaned_text = re.sub(r'\s+', ' ', text)
    return cleaned_text

def remove_non_ascii(text):
    # Use a list comprehension to filter out non-ASCII characters
    clean_text = ''.join(char for char in text if ord(char) < 128)
    return clean_text

def example_remove_non_ascii():
    text = "This is a text with non-ASCII characters: Café, rôle, façade."
    cleaned_text = remove_non_ascii(text)
    print(cleaned_text)
    
    # To convert to ascii json:
    # json_data = json.dumps({"text": cleaned_text}, ensure_ascii=False)
    # print(json_data)
    
def clean_the_text(text):
    clean_nonprintables = remove_non_printable_except_newline(text) # problem: this also removes new line chars
    clean_spaces = clean_consecutive_spaces(clean_nonprintables)
    cleaned_text = remove_non_ascii(clean_spaces)
    return cleaned_text

def extract_metadata_for_docx(file_path: str):
    """Extract metadata for a single file (file_name, author, title, date_created, date_last_updated, file_type, keywords, comments)

    Args:
        file_path (str): File path and name

    Returns:
        dict: Dictionary containing document properties
    """
    # Initialize the result dictionary
    metadata = {
        'file_name': os.path.basename(file_path),
        'author': 'Donn Pienaar',
        'title': None,
        'date_created': None,
        'date_last_updated': None,
        'file_type': os.path.splitext(file_path)[1][1:].lower(),  # Get the file extension as the file type
        'keywords': None,
        'comments': None
    }

    # if document_info["file_type"] == "pdf":
    #     try:
    #         with open(file_path, "rb") as pdf_file:
    #             pdf_reader = PyPDF2.PdfFileReader(pdf_file)
    #             info = pdf_reader.getDocumentInfo()
    #             document_info["author"] = info.author
    #             document_info["title"] = info.title
    #             document_info["date_created"] = info.created
    #             document_info["date_last_updated"] = info.modDate
    #     except Exception as e:
    #         print(f"An error occurred while extracting PDF document info: {e}")
    if metadata["file_type"] == "docx":
        try:
            doc = Document(file_path)
            metadata["author"] = doc.core_properties.author
            metadata["title"] = doc.core_properties.title
            metadata["date_created"] = doc.core_properties.created
            metadata["date_last_updated"] = doc.core_properties.modified
            metadata["keywords"] = doc.core_properties.keywords
            metadata["comments"] = doc.core_properties.comments
        except Exception as e:
            print(f"An error occurred while extracting Word document info: {e}")

    return metadata

def example_extract_metadata():
    # Example usage:
    # file_path_pdf = "example.pdf"  # Replace with the path to your PDF document
    file_path_docx = "C:\\Users\donnp\\OneDrive\\8. Poly\How far does our responsibility stretch when considering metamours.docx"
    # info_pdf = extract_metadata(file_path_pdf)
    info_docx = extract_metadata_for_docx(file_path_docx)
    # print("PDF Document Info:")
    # print(info_pdf)
    print("\nWord Document Info:")
    print(info_docx)


def get_docs_and_clean(docx_start_directory):
    documents = []
    docx_files = find_docx_files(docx_start_directory)
    
    for file_path in docx_files:
        text = get_docx_text(file_path)
        cleaned_text = clean_the_text(text)
        documents.append(cleaned_text)
        
    return documents


def create_and_set_docx_properties(file_path_and_name: str, author: str, title: str, keywords: str, comments: str):
    try:
        # Create a new document
        doc = Document()
        
        # Get the core properties
        core_props = doc.core_properties
        # print(core_props)
        
        # Set the properties
        core_props.author = author
        core_props.title = title
        core_props.created = datetime.now()
        core_props.keywords = keywords
        core_props.comments = comments
        
        # if date_created:
        #     core_props.created = datetime.now()
        
        # Save the new document
        doc.save(file_path_and_name)
        
        return True

    except Exception as e:
        print(f"An error occurred while creating and setting properties: {e}")
        return False

def example_create_and_set_docx_properties():
    file_path = "new_document.docx"  # Replace with the path where you want to create the new .docx file
    title = "New custom Document Title"
    author = "Donn Pienaar"
    keywords = "kw1, kw2, key word 3" # Keywords are stored in the tags field
    comments = "This is comments"

    if create_and_set_docx_properties(file_path, author, title, keywords, comments):
        print(f"New document '{file_path}' created and properties set successfully.")
    else:
        print("Failed to create the new document and set properties.")




# Usage example:
if __name__ == "__main__":
    
    # example_find_docx_files()
    # example_get_docx_text()
    
    # file_path = "C:\\Users\donnp\\OneDrive\\8. Poly\How far does our responsibility stretch when considering metamours.docx"
    # text = get_docx_text(file_path)
    # # print(text)
    # cleaned_text = clean_the_text(text)    
    # print(cleaned_text)
    
    # example_create_and_set_docx_properties()
    
    file_path_docx = "C:\\Users\donnp\\OneDrive\\8. Poly\How far does our responsibility stretch when considering metamours.docx"
    # file_path_docx = "new_document.docx"
    info_docx = extract_metadata_for_docx(file_path_docx)
    print("\nWord Document Info:")
    print(info_docx)


    # example_detect_encoding()
    # example_find_non_displayable_characters()
    # example_remove_non_ascii()